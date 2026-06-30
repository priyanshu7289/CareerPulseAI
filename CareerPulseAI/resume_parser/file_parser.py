"""
resume_parser/file_parser.py
Extracts raw text from uploaded resume files (PDF or DOCX) so it can be
passed into skill_extractor.SkillExtractor.extract().
"""

import io
import logging
from typing import Union

import PyPDF2
import docx

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)


def extract_text_from_pdf(file: Union[str, io.BytesIO]) -> str:
    text = []
    try:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    except Exception as e:
        logger.error(f"Failed to parse PDF: {e}")
    return "\n".join(text)


def extract_text_from_docx(file: Union[str, io.BytesIO]) -> str:
    text = []
    try:
        document = docx.Document(file)
        for para in document.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
    return "\n".join(text)


def extract_resume_text(file, filename: str) -> str:
    """
    Dispatch based on file extension. `file` can be a path string or an
    in-memory file-like object (e.g. from st.file_uploader).
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file)
    elif ext == "docx":
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Unsupported resume format: .{ext} (use PDF or DOCX)")


if __name__ == "__main__":
    # Example usage with a local file path
    sample_path = "sample_resume.pdf"
    try:
        text = extract_resume_text(sample_path, sample_path)
        print(text[:500])
    except FileNotFoundError:
        print("Place a sample_resume.pdf in this directory to test extraction.")
